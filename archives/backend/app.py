import docx
from io import BytesIO
from flask import Flask, jsonify, request, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_cors import CORS
from reportlab.lib.pagesizes import letter, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from docx.shared import Inches
from scheduler import auto_schedule_courses 
from algoclass import Room as AlgoRoom, Course as AlgoCourse, TimeSlot as AlgoTimeSlot


app = Flask(__name__)
CORS(app)  # Enable CORS for frontend interaction

# Database Configuration (Ensure the path is correct)
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///unischedule.db' 
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)

# ---- DATABASE MODELS ---- #
class Lecturer(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    department = db.Column(db.String(100), nullable=False)
    courses = db.relationship('Course', backref='lecturer', lazy=True)

class Room(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    capacity = db.Column(db.Integer, nullable=False)

class Course(db.Model):
    id = db.Column(db.String(10), primary_key=True)  # Course Code
    name = db.Column(db.String(100), nullable=False)
    level = db.Column(db.Integer, nullable=False)
    num_students = db.Column(db.Integer, nullable=False)
    lecturer_id = db.Column(db.Integer, db.ForeignKey('lecturer.id'), nullable=False)
    time_slots = db.relationship('TimeSlot', backref='course', lazy=True)

class TimeSlot(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    course_id = db.Column(db.String(10), db.ForeignKey('course.id'), nullable=False)
    day = db.Column(db.String(10), nullable=False)
    start_time = db.Column(db.String(5), nullable=False)  # HH:MM format
    end_time = db.Column(db.String(5), nullable=False)

# ---- API ROUTES ---- #
@app.route('/')
def home():
    return jsonify({'message': 'Welcome to UniSchedul API'}), 200


# API Endpoints

# GET /api/lecturers
@app.route('/api/lecturers', methods=['GET'])
def get_lecturers():
    lecturers = Lecturer.query.all()
    return jsonify([
        {'id': l.id, 'name': l.name, 'department': l.department}
        for l in lecturers
    ])

# POST /api/lecturers
@app.route('/api/lecturers', methods=['POST'])
def add_lecturer():
    data = request.json
    new_lecturer = Lecturer(name=data['name'], department=data['department'])
    db.session.add(new_lecturer)
    db.session.commit()
    return jsonify({"message": "Lecturer added successfully"}), 201

# PUT /api/lecturers/<id>
@app.route('/api/lecturers/<int:id>', methods=['PUT'])
def update_lecturer(id):
    data = request.json
    lecturer = Lecturer.query.get(id)
    if lecturer:
        lecturer.name = data['name']
        lecturer.department = data['department']
        db.session.commit()
        return jsonify({"message": "Lecturer updated successfully"})
    return jsonify({"error": "Lecturer not found"}), 404

# DELETE /api/lecturers/<id>
@app.route('/api/lecturers/<int:id>', methods=['DELETE'])
def delete_lecturer(id):
    lecturer = Lecturer.query.get(id)
    if lecturer:
        db.session.delete(lecturer)
        db.session.commit()
        return jsonify({"message": "Lecturer deleted successfully"})
    return jsonify({"error": "Lecturer not found"}), 404

# --- API ROUTES FOR ROOM MANAGEMENT ---

# Fetch all rooms
@app.route('/api/rooms', methods=['GET'])
def get_rooms():
    rooms = Room.query.all()
    return jsonify([{"id": r.id, "name": r.name, "capacity": r.capacity} for r in rooms])

# Add a new room
@app.route('/api/rooms', methods=['POST'])
def add_room():
    data = request.json
    new_room = Room(name=data['name'], capacity=data['capacity'])
    db.session.add(new_room)
    db.session.commit()
    return jsonify({"message": "Room added successfully!"}), 201

# Edit an existing room
@app.route('/api/rooms/<int:room_id>', methods=['PUT'])
def update_room(room_id):
    room = Room.query.get(room_id)
    if not room:
        return jsonify({"error": "Room not found"}), 404
    
    data = request.json
    room.name = data['name']
    room.capacity = data['capacity']
    db.session.commit()
    return jsonify({"message": "Room updated successfully!"})

# Delete a room
@app.route('/api/rooms/<int:room_id>', methods=['DELETE'])
def delete_room(room_id):
    room = Room.query.get(room_id)
    if not room:
        return jsonify({"error": "Room not found"}), 404
    
    db.session.delete(room)
    db.session.commit()
    return jsonify({"message": "Room deleted successfully!"})


# --- API ROUTES FOR COURSE MANAGEMENT ---

# Fetch all courses
@app.route('/api/courses', methods=['GET'])
def get_courses():
    courses = Course.query.all()
    return jsonify([
        {
            "id": c.id,
            "name": c.name,
            "level": c.level,
            "num_students": c.num_students,
            "lecturer_id": c.lecturer_id,
            "time_slots": [{"day": ts.day, "start_time": ts.start_time, "end_time": ts.end_time} for ts in c.time_slots]
        }
        for c in courses
    ])


# Add a new course
@app.route('/api/courses', methods=['POST'])
def add_course():
    data = request.json

    # Ensure 'time_slots' is always a list, even if missing
    time_slots = data.get('time_slots', [])

    new_course = Course(
        id=data['id'], 
        name=data['name'], 
        level=data['level'], 
        num_students=data['num_students'], 
        lecturer_id=data['lecturer_id']
    )
    db.session.add(new_course)

    # Only add time slots if they exist
    for slot in time_slots:
        new_slot = TimeSlot(
            course_id=data['id'], 
            day=slot['day'], 
            start_time=slot['start_time'], 
            end_time=slot['end_time']
        )
        db.session.add(new_slot)

    db.session.commit()
    return jsonify({"message": "Course added successfully!"}), 201


# Edit a course
@app.route('/api/courses/<string:course_id>', methods=['PUT'])
def update_course(course_id):
    course = Course.query.get(course_id)
    if not course:
        return jsonify({"error": "Course not found"}), 404
    
    data = request.json
    course.name = data['name']
    course.level = data['level']
    course.num_students = data['num_students']
    course.lecturer_id = data['lecturer_id']

    # Update time slots
    TimeSlot.query.filter_by(course_id=course_id).delete()
    for slot in data['time_slots']:
        new_slot = TimeSlot(course_id=course_id, day=slot['day'], start_time=slot['start_time'], end_time=slot['end_time'])
        db.session.add(new_slot)
    
    db.session.commit()
    return jsonify({"message": "Course updated successfully!"})

# Delete a course
@app.route('/api/courses/<string:course_id>', methods=['DELETE'])
def delete_course(course_id):
    course = Course.query.get(course_id)
    if not course:
        return jsonify({"error": "Course not found"}), 404

    db.session.delete(course)
    db.session.commit()
    return jsonify({"message": "Course deleted successfully!"})


# --- API ROUTES FOR TIME SLOT MANAGEMENT ---

# Fetch all time slots
@app.route('/api/timeslots', methods=['GET'])
def get_time_slots():
    time_slots = TimeSlot.query.all()
    return jsonify([
        {
            "id": ts.id,
            "course_id": ts.course_id,
            "course_code": ts.course.id,  # Include course code
            "course_name": ts.course.name,  # Include course title
            "lecturer_name": ts.course.lecturer.name, 
            "day": ts.day,
            "start_time": ts.start_time,
            "end_time": ts.end_time
        }
        for ts in time_slots
    ])


# Add a new time slot
@app.route('/api/timeslots', methods=['POST'])
def add_timeslot():
    try:
        data = request.json
        course_id = data.get("course_id")
        
        # Validate that the course exists
        course = Course.query.get(course_id)
        if not course:
            return jsonify({"error": "Course not found"}), 404
        
        new_slot = TimeSlot(
            course_id=course_id,
            day=data["day"],
            start_time=data["start_time"],
            end_time=data["end_time"]
        )
        db.session.add(new_slot)
        db.session.commit()

        return jsonify({"message": "Time slot added successfully!", "timeslot_id": new_slot.id}), 201

    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

# Edit a time slot
@app.route('/api/timeslots/<int:timeslot_id>', methods=['PUT'])
def update_timeslot(timeslot_id):
    try:
        data = request.json
        timeslot = TimeSlot.query.get(timeslot_id)
        
        if not timeslot:
            return jsonify({"error": "Time slot not found"}), 404

        # Update the time slot details
        timeslot.day = data["day"]
        timeslot.start_time = data["start_time"]
        timeslot.end_time = data["end_time"]

        db.session.commit()
        return jsonify({"message": "Time slot updated successfully!"})

    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@app.route('/api/timeslots/<int:timeslot_id>', methods=['DELETE'])
def delete_timeslot(timeslot_id):
    try:
        timeslot = TimeSlot.query.get(timeslot_id)
        
        if not timeslot:
            return jsonify({"error": "Time slot not found"}), 404

        db.session.delete(timeslot)
        db.session.commit()
        return jsonify({"message": "Time slot deleted successfully!"})

    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

# --- API ROUTES FOR TIMETABLE MANAGEMENT ---
@app.route('/api/timetable', methods=['GET'])
def get_timetable():
    """Fetch all scheduled courses organized by days."""
    week_days = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday"]
    timetable = {day: [] for day in week_days}

    courses = Course.query.all()
    for course in courses:
        for slot in course.time_slots:
            timetable[slot.day].append({
                "course_code": course.id,
                "course_name": course.name,
                "lecturer": course.lecturer.name,
                "start_time": slot.start_time,
                "end_time": slot.end_time,
                "room": "N/A"  # We can enhance this if room allocation is added
            })

    return jsonify(timetable)


# API Algorithm Endpoint
@app.route('/api/run-algorithm', methods=['GET'])
def run_algorithm():
    """Fetches timetable data, runs the scheduling algorithm, and returns results"""
    
    # Fetch all necessary data from the database
    rooms = Room.query.all()
    courses = Course.query.all()

    # Convert data into algorithm-readable format
    room_list = [AlgoRoom(r.id, r.name, r.capacity) for r in rooms]
    course_list = [
        AlgoCourse(
            c.id, 
            c.name, 
            c.level, 
            c.num_students, 
            [AlgoTimeSlot(ts.day, ts.start_time, ts.end_time) for ts in c.time_slots], 
            c.lecturer_id
        ) for c in courses
    ]

    # Initialize logs
    logs = []  # ✅ Add this so logs can be recorded properly

    # Run the scheduling algorithm with all required arguments
    bookings, failed_bookings = auto_schedule_courses(course_list, room_list, logs)

    # Fetch lecturer names for lookup
    lecturer_lookup = {lec.id: lec.name for lec in Lecturer.query.all()}  # ✅ Create a mapping of lecturer ID -> Name

    # Format results for JSON response
    result = {
        "bookings": [  # ✅ Scheduled timetable data
            {
                "course_id": b.course.course_id,
                "course_name": b.course.name,
                "lecturer": lecturer_lookup.get(b.course.lecturer_id, "Unknown Lecturer"),  # ✅ Get lecturer name
                "room": b.room.name,
                "day": b.day,
                "start_time": b.start_time,
                "end_time": b.end_time
            }
            for b in bookings
        ],
        "failed_bookings": failed_bookings,  # ✅ Failed scheduling attempts
        "logs": logs  # ✅ Include logs in the response
    }

    return jsonify(result), 200

# Function to structure the schedule like OptimizedSchedule.jsx
def format_schedule_for_pdf(bookings, failed_bookings):
    week_days = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday"]

    # Sort bookings first by day, then by time
    sorted_bookings = sorted(bookings, key=lambda x: (week_days.index(x["day"]), x["start_time"]))

    # Organize bookings by day
    schedule_by_day = {day: [] for day in week_days}
    for booking in sorted_bookings:
        schedule_by_day[booking["day"]].append(
            f"🔹 {booking['start_time']} - {booking['end_time']}: {booking['course_name']} "
            f"in {booking['room']} by Lecturer {booking.get('lecturer', 'Unknown')}"
        )

    # Organize failed bookings
    formatted_failed_bookings = [f"❌ {failure}" for failure in failed_bookings]

    return schedule_by_day, formatted_failed_bookings

@app.route('/api/dashboard-stats', methods=['GET'])
def get_dashboard_stats():
    """Fetch summary statistics for the dashboard."""
    return jsonify({
        "courses": Course.query.count(),
        "lecturers": Lecturer.query.count(),
        "rooms": Room.query.count(),
        "timeslots": TimeSlot.query.count(),
    })

# @app.route("/api/generate-pdf", methods=["POST"])
# def generate_pdf():
    data = request.json

    # Extract scheduling data
    schedule = data.get("schedule", [])
    failed = data.get("failed_bookings", [])
    semester = data.get("semester", "Semester")
    academic_year = data.get("academic_year", "Year")
    department = data.get("department", "Department")

    # 1) Create PDF doc in memory, using LANDSCAPE
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=landscape(letter),    # <-- ensures a wide orientation
        leftMargin=20,
        rightMargin=20,
        topMargin=20,
        bottomMargin=20
    )

    story = []
    styles = getSampleStyleSheet()

    # 2) Title
    title_par = Paragraph(f"<b>{department} - {semester} ({academic_year})</b>", styles["Title"])
    story.append(title_par)
    story.append(Spacer(1, 12))

    # 3) Build a table for your timetable
    DAYS = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday"]
    HOURS = [
        "08:00","09:00","10:00","11:00","12:00",
        "13:00","14:00","15:00","16:00","17:00"
    ]

    # Table header (row 0)
    table_data = [["Days"] + HOURS]

    # fill each row with day + cells for each hour
    for day in DAYS:
        row = [day]
        for hour in HOURS:
            # gather all classes in [hour, hour+1)
            booked = []
            for b in schedule:
                if b["day"] == day and hour >= b["start_time"] and hour < b["end_time"]:
                    text = f"{b['course_id']} - {b['course_name']}\n{b['lecturer']}\nRoom: {b['room']}"
                    booked.append(text)
            row.append("\n\n".join(booked) if booked else "")
        table_data.append(row)

    tbl = Table(table_data, repeatRows=1)
    # 4) Style the table
    tbl_style = TableStyle([
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('BACKGROUND', (0, 0), (-1, 0), colors.gray),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 11),
        ('BACKGROUND', (0, 1), (0, -1), colors.lightgrey),
        ('ALIGN', (0, 1), (0, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),   
        ('FONTSIZE', (0, 1), (-1, -1), 5),    
        ('LEFTPADDING', (0, 1), (-1, -1), 3),
        ('RIGHTPADDING', (0, 1), (-1, -1), 3),
        ('TOPPADDING', (0, 1), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 3),
    ])
    tbl.setStyle(tbl_style)
    story.append(tbl)
    story.append(Spacer(1, 12))

    # 5) Display failed bookings if any
    if failed:
        story.append(Paragraph("Failed Scheduling Attempts:", styles["Heading2"]))
        fails_data = []
        for f in failed:
            fails_data.append([f])
        fail_table = Table(fails_data, colWidths=[700])
        fail_table.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.black),
            ('BACKGROUND', (0,0), (-1,0), colors.red),
            ('TEXTCOLOR', (0,0), (-1,0), colors.white),
            ('FONTNAME', (0,0), (-1,-1), 'Helvetica'),
            ('FONTSIZE', (0,0), (-1,-1), 6),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ]))
        story.append(fail_table)

    # Build & return
    doc.build(story)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="Optimized_Schedule.pdf", mimetype="application/pdf")

@app.route("/api/export-file", methods=["POST"])
def export_file():
    data = request.json
    file_format = data.get("format", "pdf")  # "pdf" or "docx"

    # Extract the fields
    schedule = data.get("schedule", [])
    failed   = data.get("failed_bookings", [])
    semester = data.get("semester", "Semester")
    faculty  = data.get("faculty", "Faculty")
    session  = data.get("session", "Session")
    year     = data.get("academic_year", "Year")
    dept     = data.get("department", "Department")

    if file_format == "pdf":
        return generate_pdf(schedule, failed, semester, year, dept, faculty, session)
    else:
        return generate_docx(schedule, failed, semester, year, dept, faculty, session)

def generate_pdf(schedule, failed, semester, year, dept, faculty, session):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(letter))

    styles = getSampleStyleSheet()
    story = []

    # Title
    title_par = Paragraph(f"<b>{dept}, {faculty} - {semester} ({year})<br/>{session}</b>", styles["Title"])
    story.append(title_par)
    story.append(Spacer(1, 12))

    # Build a table
    DAYS = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday"]
    HOURS = ["08:00","09:00","10:00","11:00","12:00","13:00","14:00","15:00","16:00","17:00"]
    table_data = [["Days"] + HOURS]
    for day in DAYS:
        row = [day]
        for hour in HOURS:
            booked = []
            for b in schedule:
                if b["day"] == day and hour >= b["start_time"] and hour < b["end_time"]:
                    txt = f"{b['course_id']} - {b['course_name']}\n{b['lecturer']}\nRoom: {b['room']}"
                    booked.append(txt)
            row.append("\n\n".join(booked) if booked else "")
        table_data.append(row)

    tbl = Table(table_data, repeatRows=1)
    style = TableStyle([
        ("GRID", (0,0), (-1,-1), 1, colors.black),
        ("BACKGROUND", (0,0), (-1,0), colors.gray),
        ("TEXTCOLOR", (0,0), (-1,0), colors.whitesmoke),
        ("ALIGN", (0,0), (-1,0), "CENTER"),
        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTSIZE", (0,0), (-1,0), 11),
        ("BACKGROUND", (0,1), (0,-1), colors.lightgrey),
        ("ALIGN", (0,1), (0,-1), "CENTER"),
        ("VALIGN", (0,0), (-1,-1), "TOP"),
        ("FONTSIZE", (0,1), (-1,-1), 9),
    ])
    tbl.setStyle(style)
    story.append(tbl)

    # Failed
    if failed:
        story.append(Spacer(1, 12))
        story.append(Paragraph("Failed Scheduling Attempts:", styles["Heading2"]))
        fails_data = []
        for f in failed:
            fails_data.append([f])
        fail_table = Table(fails_data, colWidths=[700])
        fail_table.setStyle(TableStyle([
            ('GRID',(0,0),(-1,-1),0.5,colors.black),
            ('BACKGROUND',(0,0),(-1,0),colors.red),
            ('TEXTCOLOR',(0,0),(-1,0),colors.white),
            ('FONTNAME',(0,0),(-1,-1),'Helvetica'),
            ('FONTSIZE',(0,0),(-1,-1),9),
            ('VALIGN',(0,0),(-1,-1),'TOP'),
        ]))
        story.append(fail_table)

    doc.build(story)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True,
                     download_name="Optimized_Schedule.pdf",
                     mimetype="application/pdf")

def generate_docx(schedule, failed, semester, year, dept, faculty, session):
    # We'll build a .docx in memory
    doc = docx.Document()

    doc.add_heading(f"{dept}, {faculty} - {semester} ({year}) / {session}", 0)

    DAYS = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday"]
    HOURS = ["08:00","09:00","10:00","11:00","12:00",
             "13:00","14:00","15:00","16:00","17:00"]

    # Create table with rows = 1 + len(DAYS), columns = 1 + len(HOURS)
    rows = 1 + len(DAYS)
    cols = 1 + len(HOURS)
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'

    # First row: hour headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Days"
    for i, hour in enumerate(HOURS):
        hdr_cells[i+1].text = hour

    # Fill day labels in first column
    for r, day in enumerate(DAYS, start=1):
        day_cell = table.cell(r, 0)
        day_cell.text = day

    # Fill in each cell
    for r, day in enumerate(DAYS, start=1):
        for c, hour in enumerate(HOURS, start=1):
            found = []
            for b in schedule:
                if b["day"] == day and hour >= b["start_time"] and hour < b["end_time"]:
                    text = f"{b['course_id']} - {b['course_name']}\n{b['lecturer']}\nRoom: {b['room']}"
                    found.append(text)
            cell_text = "\n\n".join(found) if found else ""
            table.cell(r, c).text = cell_text

    if failed:
        doc.add_heading("Failed Scheduling Attempts:", level=2)
        for f in failed:
            doc.add_paragraph(f, style='List Bullet')

    # Save to memory
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True,
                     download_name="Optimized_Schedule.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.route('/api/recent-logs', methods=['GET'])
def get_recent_logs():
    """Fetch recent scheduling logs."""
    logs = [
        "📌 Course PHY101 scheduled on Monday 08:00 - 10:00",
        "⚠️ Conflict: Lecturer Dr. Smith assigned to two courses at the same time!",
        "📌 New course CSC202 added to the database",
    ]
    return jsonify(logs)


@app.route("/api/lecturers/<int:lecturer_id>/courses", methods=['GET'])
def get_lecturer_courses(lecturer_id):
    """Fetch a lecturer + all assigned courses."""
    lecturer = Lecturer.query.get(lecturer_id)
    if not lecturer:
        return jsonify({"error": "Lecturer not found"}), 404
    
    # Build response
    courses_data = []
    for c in lecturer.courses:
        courses_data.append({
            "id": c.id,
            "name": c.name,
            "level": c.level,
            "num_students": c.num_students
        })

    return jsonify({
        "id": lecturer.id,
        "name": lecturer.name,
        "department": lecturer.department,
        "courses": courses_data
    })

@app.route("/api/lecturers/<int:lecturer_id>/courses", methods=['POST'])
def add_course_to_lecturer(lecturer_id):
    """Assign an existing or new course to this lecturer."""
    data = request.json
    lecturer = Lecturer.query.get(lecturer_id)
    if not lecturer:
        return jsonify({"error": "Lecturer not found"}), 404

    # We expect { "course_id": "...", "name": "...", "level": ..., "num_students": ... }
    #  If course_id already exists, we reassign the course's lecturer_id
    #  Otherwise, create new course
    course_id = data["course_id"]
    course = Course.query.get(course_id)
    if course:
        # Reassign
        course.lecturer_id = lecturer_id
        db.session.commit()
        return jsonify({"message": f"Course {course_id} assigned to lecturer {lecturer_id}"}), 200
    else:
        # Create new
        new_course = Course(
            id=course_id,
            name=data["name"],
            level=data["level"],
            num_students=data["num_students"],
            lecturer_id=lecturer_id
        )
        db.session.add(new_course)
        db.session.commit()
        return jsonify({"message": f"New course {course_id} created & assigned"}), 201

@app.route("/api/lecturers/<int:lecturer_id>/courses/<string:course_id>", methods=['PUT'])
def update_lecturer_course(lecturer_id, course_id):
    """Update details of a course belonging to this lecturer."""
    data = request.json
    course = Course.query.filter_by(id=course_id, lecturer_id=lecturer_id).first()
    if not course:
        return jsonify({"error": "Course not found for this lecturer"}), 404
    
    # Update fields
    course.name = data.get("name", course.name)
    course.level = data.get("level", course.level)
    course.num_students = data.get("num_students", course.num_students)
    db.session.commit()
    return jsonify({"message": f"Course {course_id} updated"}), 200

@app.route("/api/lecturers/<int:lecturer_id>/courses/<string:course_id>", methods=['DELETE'])
def remove_course_from_lecturer(lecturer_id, course_id):
    """Remove a course from this lecturer (set lecturer_id to None or reassign) or delete the course."""
    course = Course.query.filter_by(id=course_id, lecturer_id=lecturer_id).first()
    if not course:
        return jsonify({"error": "Course not found or not assigned to this lecturer"}), 404
    
    # Option A: DELETE the course entirely
    db.session.delete(course)
    db.session.commit()
    return jsonify({"message": f"Course {course_id} removed from lecturer {lecturer_id} and deleted."}), 200


if __name__ == '__main__':
    with app.app_context():
        db.create_all()  # Ensure the database and tables exist
    app.run(debug=True)

